#! python3
# Your one stop destination for all your assignment needs 
# Convert your assignments to PDFs, watermark them, merge them and encrypt them, all in one place
# Get to know the birthdays of your classmates and send them a birthday wish to show you care
# Convert your PDF into audiobook and learn on the go


#We import all the required modules required for the program
import os, openpyxl, PyPDF2, pyttsx3, docx, logging, sys, time
from openpyxl.styles import Font, Color, NamedStyle, Alignment, Border, Side, colors
from docx2pdf import convert
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime #retrieve date and time
from termcolor import colored #For adding colored text

#Function to clear the output screen
def screen_clear():
    _ = os.system('cls')


'''
    birthdays.xlsx has the names and birth days of all students in the class
    We pass "birthdays" as the filename.
    Here we retrieve the names of those having birthdays in the current month.
    We can also find out the birthday of any person present in the class
'''

def excel_operations(filename):
    screen_clear()
    if not os.path.exists(filename + ".xlsx"):
        logging.error(colored('File does not exist', 'red')) # if file does not exist, return the control
        return 

    screen_clear() # clear the screen
    wb = openpyxl.load_workbook(filename + '.xlsx')
    # takes in the filename and returns a value of the workbook data type
    sheet = wb['Sheet1'] # Worksheet object

    #Freeze the first row so that it is 
    #always visible to the user even as they scroll through the spreadsheet
    sheet.freeze_panes = 'A2'

    #Set font styles to the first row
    boldFont = Font(bold=True)
    center_align = Alignment(horizontal="center", vertical="center")
    border_text = Border(bottom=Side(border_style="thin"))

    header_row = sheet[1]
    for cell in header_row:
         cell.font = boldFont
         cell.alignment = center_align
         cell.border = border_text


    names = [] # To store names of those having birthdays in the current month
    bday = {} # To store the birthdays of the class people

    count = 0
    for row in range(2, sheet.max_row + 1):
        # Retrieve values of the cells from the sheet
        name = sheet['A' + str(row)].value
        day = sheet['B' + str(row)].value
        month = sheet['C' + str(row)].value
        year = sheet['D' + str(row)].value

        #Create the birthday of each person in dd-mm-yyyy format
        birthday = str(day) + '-' + month + '-' + str(year)
        #Store the birthday in the bday dictionary
        bday[name] = birthday

        currentMonth = datetime.now().strftime('%B') #Obtain the current month, say December

        # If there is a birthday in the current month, set the value of the cell to True, else False
        if currentMonth == month:
            sheet['E' + str(row)] = 'True'
            names.append(name)
            count += 1
        else:
            sheet['E' + str(row)] = 'False'

    # Display the name of the person having birthday in the current month
    if count == 0:
        print("There aren't any birthdays in this month") 

    if count == 1:
        print("There is " + str(count) + " person with birthday in this month") 
        for name in names:
            print(colored(name, "green")) 

    else:
        print("There are " + str(count) + " people with birthdays in this month") 
        for name in names:
            print(colored(name, "green"))  

    # Users can retrive the birthdays of other people as well
    print("Do you want to find out someone's birthday? (Enter 1 if YES)")

    try:
        choice = int(input())
    except:
        logging.warning(colored('You should have entered a number', 'red')) # if input is not an integer
        print("\n")
        return

    if choice == 1:
        
        print("Enter the name")
        person = input()
        person = person.title() # To get the names in the format it is stored in the spreadsheet

        try:
            print(person + "'s birthday is on " + str(bday[person]))
        except:
            print(person + " isn't part of your class")
    
    wb.save('birthdays.xlsx') #We save our changes done in the birthdays.xlsx file
    print('\n' * 2)



'''
    We can create a birthday card for any person.
    The card gets stored in the birthday_wish.docx file.
    You can then print this card and send it to your friends
'''

def birthday_card():
    screen_clear() #clear the screen

    doc = docx.Document()
    print("Enter the name of the birthday girl/boy")
    name = input()
    name = name.title()
    
    p = doc.add_paragraph("Happy Birthday " + name, "Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER #Align the birthday wish at the center of the document
    

    doc.add_picture('bday.png',  width=docx.shared.Inches(6), height=docx.shared.Cm(12)) #Add a birthday card
    doc.add_page_break() 
        
    doc.save('birthday_wish.docx') # Save the birthday card in birthday_wish.docx file
    print(colored('Birthday Card has been generated', 'blue'))
    print('\n' * 2)



'''
    Customise a class schedule for every person in the class.
    All you have to do is enter the sem, section and names of the students.
    The schedules can be found in the schedules.docx file in the current working directory.
    There is a schedule for each student on a new page, so you can open a single Word document
    and print all schedules together.
'''

def schedule_generator():
    screen_clear() #clear the screen

    doc = docx.Document() # Returns a Document object
    print("Enter the semester")
    sem = input()
    print("Enter the section")
    section = input()

    names = [] # To store names of the students
    print("List the names of the students")
    while True:
        print("Enter name of the student" + str(len(names) + 1) + " or enter nothing to stop")
        name = input()
        if name == '':
            break
        names.append(name) # Add the names to the list
    
    # Create a schedule for each student in the schedules.docx file
    for i in range(len(names)):
        p = doc.add_paragraph(names[i] + "'s Class Schedule", "Title")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph("Sem: " + sem + "       Section: " + section, "Normal")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_picture('schedule.png',  width=docx.shared.Inches(6), height=docx.shared.Cm(12))
        doc.add_page_break()
        
    print(colored('Schedules are completed', 'blue'))
    doc.save('schedules.docx') # save to schedules.docx file
    print('\n' * 2)



'''
    It is always better to watermark your documents and assignments with your USN or name.
    All you have to do is save your watermarked PDF as watermark.pdf in the current working directory.
    You can then pass in any PDf file into the program and apply your watermark to the desired pages of your document.
    
    You can pass "hackers" as the filename. Specify the starting and ending page numbers, say 2 and 5. 
    You can give the filename where you want your watermarked PDF stored.

'''

def pdf_watermark(filename):
    screen_clear()  # clear the screen

    if not os.path.exists(filename + ".pdf"):
        logging.error(colored('File does not exist', 'red'))
        return 


    pdfFile = open(filename + '.pdf', 'rb') # open the file in read binary mode
    pdfReader = PyPDF2.PdfFileReader(pdfFile) # pdfFileReader object represents the PDf file

    print("Enter the page you want to start the watermark from")
    
    try:
        page_start  = int(input())
    except:
        logging.warning(colored('You should have entered a number', 'red')) # if input is not a number
        print("\n")
        return

    print("Enter the page you want to end the watermark")
    
    try:
        page_end  = int(input())
    except:
        logging.warning(colored('You should have entered a number', 'red')) # if input is not a number
        print("\n")
        return
    
    # If user enters a number greater than the pages present
    if page_end > pdfReader.numPages: 
        logging.warning(colored('You have exceeded the total number of pages present', 'yellow'))
        return

    print("What do you want your new filename to be? (Avoid existing filenames)")
    userfilename = input()
 
    pdfWatermarkReader = PyPDF2.PdfFileReader(open('watermark.pdf', 'rb')) # open the watermark.pdf file
    
    pdfWriter = PyPDF2.PdfFileWriter() # PDFFileWriter object
    
    # We loop through the pages can watermark the pages specified by the users
    # We then loop through the remaining pages and add the pages to the new file
    for pageNum in range(page_start - 1):
        pageObj = pdfReader.getPage(pageNum)
        pdfWriter.addPage(pageObj)

    for pageNum in range(page_start-1, page_end):
        pageObj = pdfReader.getPage(pageNum)
        pageObj.mergePage(pdfWatermarkReader.getPage(0))
        pdfWriter.addPage(pageObj)


    for pageNum in range(page_end, pdfReader.numPages):
        pageObj = pdfReader.getPage(pageNum)
        pdfWriter.addPage(pageObj)

    resultPdfFile = open(userfilename+'.pdf', 'wb') # Store the watermarked file in the filename given for the new file
    pdfWriter.write(resultPdfFile)
    pdfFile.close()
    resultPdfFile.close()
    print(colored('PDF file has been watermarked', 'blue'))
    print('\n' * 2)


'''
    Convert your Word assignments into PDF.
    You can find the PDf file in your current working directory
'''
def word_converter(filename):
    screen_clear()
    if not os.path.exists(filename + ".docx"):
        logging.error(colored('File does not exist', 'red')) # If file does not exist
        return 

    # Convert the Word document to PDF file
    convert(filename+".docx")
    convert(filename + ".docx", filename + ".pdf")
    convert("/")

    print(colored('Word file is converted to PDF', 'blue'))
    print('\n' * 2)


'''
    Here we can merge the desired PDF files. This code merges all the pages of a PDF together. 
    You can pass in "harry_potter" and "manual" in the filename. 
    Specify the new filename where you want the merged PDFs to get stored.
'''

def pdf_merger():
    screen_clear()
    print("What do you want your new filename to be? (Avoid existing filenames)")
    userfilename = input()
    files = []
    # List all the files needed to be merged
    print("List the files to be merged")
    while True:
        print("Enter name of the file" + str(len(files) + 1) + " or enter nothing to stop")
        name = input()

        if name == '':
            break

        if not os.path.exists(name + ".pdf"):
            logging.error(colored('File does not exist', 'red')) # If file does not exist
            continue

        files.append(name + ".pdf")
    

    pdfWriter = PyPDF2.PdfFileWriter() # PDFFileWriter object

    for filename in files:
        pdfFileObj = open(filename, 'rb') # open each file and add the pages to the PDFFileWriter object
        pdfReader = PyPDF2.PdfFileReader(pdfFileObj)

        for pageNum in range(0, pdfReader.numPages):
            pageObj = pdfReader.getPage(pageNum)
            pdfWriter.addPage(pageObj)

    # Save the resulting PDF to a file.
    pdfOutput = open(userfilename+'.pdf', 'wb')
    pdfWriter.write(pdfOutput)
    pdfOutput.close()
    print(colored('PDFs are merged together', 'blue'))
    print('\n' * 2)


'''
    Convert your PDF into an audiobook by passing the filename.
    You can't pause the audiobook in between. You need to close the program if you wish to stop the reading. 
    However, once the reading is done, it returns back to the menu.
'''
def pdf_speak(filename):
    screen_clear()
    if not os.path.exists(filename + ".pdf"):
        logging.error(colored('File does not exist', 'red'))
        return 

    # Display warning to the user
    print(colored('Warning : ', 'yellow'), colored('You cannot pause the audiobook in between', 'yellow'))

    print("Do you want to continue?")
    answer = input()
    answer = answer.lower()

    if answer == 'yes':
        print("\n") 
        book = open(filename + '.pdf', 'rb') # Open the file to be converted into audiobook
        pdfReader = PyPDF2.PdfFileReader(book)
        pages = pdfReader.numPages # Read the total number of pages

        speaker = pyttsx3.init()
        for num in range(0, pages):
            page = pdfReader.getPage(num) # Get each page
            text = page.extractText()
            print(colored(text, "magenta")) # Display the text to the user
            speaker.say(text) # Read out the text
            speaker.runAndWait()
        speaker.stop() # Stop the reading

        print(colored('Reading is completed', 'blue'))
        print('\n' * 2)


'''
    Secure your PDF files with an encryption and save it as a new file. 
    You can pass "harry_potter" as the filename. Give the new filename and the password to encrypt. 
    The encrypted file can be found in the filename you have given.
    Before anyone can view the encrypted PDF, they’ll have to enter the password.
'''
def pdf_encrpyt(filename):
    screen_clear() # Clear the screen
     
    if not os.path.exists(filename + ".pdf"):
        logging.error(colored('File does not exist', 'red')) # File does not exist
        return 

    pdfReader = PyPDF2.PdfFileReader(open(filename + '.pdf', 'rb'))
    pdfWriter = PyPDF2.PdfFileWriter()

    print("What do you want your new filename to be? (Avoid existing filenames)")
    userfilename = input()

    for pageNum in range(pdfReader.numPages):
        pdfWriter.addPage(pdfReader.getPage(pageNum)) # copy the pages to the PDFFileWriter object

    # Ask for encryption
    print("Enter the password")
    password = input()

    pdfWriter.encrypt(password) # Encrypt the file with the password

    resultPdf = open(userfilename+'.pdf', 'wb')
    pdfWriter.write(resultPdf)

    print(colored('Your file has been encrypted', 'blue'))
    resultPdf.close()
    print('\n' * 2)


'''
    This is to verify the PDF's password. 
    You can pass in the filename(name of the encrypted file) of the previous option to check the password 
    (or any encrypted file from the cwd).
'''
def pdf_decrpyt(filename):
    screen_clear() # clear the screen

    if not os.path.exists(filename + ".pdf"): 
        logging.error(colored('File does not exist', 'red')) # The file does not exist
        return 

    pdfReader = PyPDF2.PdfFileReader(open(filename + '.pdf', 'rb')) 

    # If file is not encrypted, return the control
    if not pdfReader.isEncrypted:
        print(colored("Your PDF file is not encrypted. Pass in an encrypted filenames", "red"))
        return

    print("Enter the password")
    password = input()

    pdfReader.decrypt(password) # Decrypt the encrypted file with the PDF

    try: # if file has been decrypted
        pageobj = pdfReader.getPage(0)
        print(colored('Your file is decrypted. The password matches', 'blue'))
    except:
        logging.warning(colored("Couldn't decrypt the file. Kindly check the password entered", "red"))
    
    print('\n' * 2)




# Menu to ask the user's choice and perform corresponding operations

choice = 1

while(choice):
    screen_clear() # clear the screen

    print(colored("-------------MENU-------------", "magenta"))
    print(colored("Enter 1 to find about birthdays", "cyan"))
    print(colored("Enter 2 to print your class schedule", "cyan"))
    print(colored("Enter 3 for watermarking your PDF", "cyan"))
    print(colored("Enter 4 to convert your Word file to PDF", "cyan"))
    print(colored("Enter 5 to merge the desired PDF files together", "cyan"))
    print(colored("Enter 6 to convert your PDF into audiobook", "cyan"))
    print(colored("Enter 7 to secure your PDF file with an encryption", "cyan"))
    print(colored("Enter 8 if you want to check your PDF password", "cyan"))
    print(colored("Enter 9 to wish your classmate on their birthday", "cyan"))
    print(colored("Enter 0 to quit", "cyan"))

    try:
        choice = int(input())
    except:
        logging.warning(colored('You should have entered a number', 'red')) # if number is not entered
        time.sleep(2)
        continue
        print("\n")
  
    print(colored("Please pass only the filename. The extensions will be added by the program.", "yellow"))

    # Perform necessary actions basis user's choice

    if choice == 1:
        print(colored("Enter the filename", "green"))
        filename = input()
        excel_operations(filename)
    if choice == 2:
        schedule_generator()
    if choice == 3:
        print(colored("Enter the filename", "green"))
        filename = input()
        pdf_watermark(filename)
    if choice == 4:
        print(colored("Enter the filename", "green"))
        filename = input()
        word_converter(filename)
    if choice == 5:
        pdf_merger()
    if choice == 6:
        print(colored("Enter the filename", "green"))
        filename = input()
        pdf_speak(filename)
    if choice == 7:
        print(colored("Enter the filename", "green"))
        filename = input()
        pdf_encrpyt(filename)
    if choice == 8:
        print(colored("Enter the filename", "green"))
        filename = input()
        pdf_decrpyt(filename)
    if choice == 9:
        birthday_card()
    if choice == 0:
        sys.exit('Thank you')

    print(colored("Do you want to continue? (Enter 1 for Yes and 0 for No)", "yellow"))
    
    try:
        option = int(input())
    except:
        logging.warning(colored('You should have entered a number', 'red'))
        print("\n")

    if option == 1:
        continue
    else:
        sys.exit("Thank you")